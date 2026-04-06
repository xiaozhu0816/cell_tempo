#!/usr/bin/env python3
"""
Fill the Cell Reports Key Resources Table (KRT) for the cell_tempo paper.
Reads the empty template, populates Table 0 with our paper's entries,
deletes the Appendix A example tables, and saves as KRT_cell_tempo_filled.docx.

Citation format: use [[N]] in strings to produce superscript reference numbers,
matching the numbered bibliography style used in main.tex.
"""

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path(__file__).parent / "Table_Template_Cell_Reports.docx"
DST = Path(__file__).parent / "KRT_cell_tempo_filled.docx"

# ── KRT entries grouped by subheading ─────────────────────────────────
# Each entry is (REAGENT or RESOURCE, SOURCE, IDENTIFIER)
# Use [[N]] for superscript citation numbers (Cell Press numbered style).
KRT = {
    "Bacterial and virus strains": [
        ("Venezuelan equine encephalitis virus (VEEV), vaccine strain TC-83",
         "P1 laboratory stock generated from an infectious cDNA clone provided by I. Frolov",
         "Kinney et al.[[19]]"),
    ],
    "Chemicals, peptides, and recombinant proteins": [
        ("Dulbecco's phosphate-buffered saline (DPBS)",
         "Corning",
         "Cat#21-030-CV"),
        ("EGM-2 Endothelial Cell Growth Medium-2 BulletKit",
         "Lonza",
         "Cat#CC-3162"),
        ("Fetal bovine serum (FBS)",
         "Gibco",
         "Cat#A5256801"),
        ("Dulbecco's Modified Eagle Medium (DMEM)",
         "Quality Biological",
         "Cat#112-014-101CS"),
        ("Trypsin (0.25%)-EDTA (0.02%)",
         "Quality Biological",
         "Cat#118-093-721"),
        ("Minimum Essential Medium (EMEM) (2X) without Phenol Red and L-Glutamine",
         "Quality Biological",
         "Cat#115-073-101"),
        ("Penicillin/Streptomycin solution (100X)",
         "Corning",
         "Cat#30-002-CI"),
    ],
    "Experimental models: Cell lines": [
        ("Human: Primary brain microvascular endothelial cells (HBMVEC), male pediatric donor, passages 5-8",
         "Cell Systems (an AnaBios company)",
         "Cat#ACBRI 376; Lot#376.07.05.01.2F"),
        ("African green monkey kidney epithelial cells (Vero)",
         "ATCC",
         "Cat#CCL-81"),
    ],
    "Software and algorithms": [
        ("CELLCYTE Studio",
         "CYTENA",
         "https://www.cytena.com/products/cellcyte-x/"),
        ("GraphPad Prism v11.0.0",
         "GraphPad Software",
         "https://www.graphpad.com"),
        ("Python (v3.10)",
         "Python Software Foundation",
         "https://www.python.org"),
        ("PyTorch (v2.0)",
         "PyTorch",
         "https://pytorch.org"),
        ("torchvision (ResNet50 with ImageNet pretrained weights)",
         "PyTorch",
         "https://pytorch.org/vision"),
        ("ResNet50 architecture",
         "He et al.[[20]]",
         "https://pytorch.org/vision/stable/models/resnet.html"),
        ("scikit-learn",
         "Pedregosa et al.[[22]]",
         "https://scikit-learn.org"),
        ("t-SNE (visualization)",
         "van der Maaten and Hinton[[21]]",
         "https://scikit-learn.org/stable/modules/generated/sklearn.manifold.TSNE.html"),
        ("cell_tempo analysis pipeline",
         "This paper",
         "https://github.com/CAIR-LAB-WFUSM/cell_tempo"),
    ],
    "Other": [
        ("CELLCYTE X live-cell imaging system (10x objective, Enhanced Contour mode)",
         "CYTENA GmbH, Freiburg, Germany",
         "CELLCYTE X Quick Start Guide; https://www.cytena.com/products/cellcyte-x/"),
        ("12-well tissue-culture plates (CELLSTAR)",
         "Greiner Bio-One",
         "Cat#665180"),
    ],
}

# Subheadings present in template (in order); ones not in KRT will be removed.
SUBHEADINGS_IN_TEMPLATE = [
    "Antibodies",
    "Bacterial and virus strains",
    "Biological samples",
    "Chemicals, peptides, and recombinant proteins",
    "Critical commercial assays",
    "Deposited data",
    "Experimental models: Cell lines",
    "Experimental models: Organisms/strains",
    "Oligonucleotides",
    "Recombinant DNA",
    "Software and algorithms",
    "Other",
]

# Regex to find [[N]] superscript markers
_SUPER_RE = re.compile(r"\[\[(\d+)\]\]")


def first_cell_text(row):
    return row.cells[0].text.strip()


def is_subheading_row(row):
    """Heading rows have all 3 cells == the same heading text."""
    txts = [c.text.strip() for c in row.cells]
    return txts[0] == txts[1] == txts[2] != "" and txts[0] != "REAGENT or RESOURCE"


def remove_row(table, row):
    """Remove a row from a python-docx table."""
    row._element.getparent().remove(row._element)


def insert_row_after(table, ref_row):
    """Insert a new empty row immediately after ref_row, return the new row."""
    new_tr = deepcopy(ref_row._element)
    for tc in new_tr.iter(qn("w:tc")):
        for p in tc.iter(qn("w:p")):
            for r in list(p):
                tag = r.tag.split("}")[-1]
                if tag == "r":
                    p.remove(r)
    ref_row._element.addnext(new_tr)
    for r in table.rows:
        if r._element is new_tr:
            return r
    return None


def set_cell_text(cell, text):
    """
    Write text into the first paragraph of cell.
    Supports [[N]] notation: those tokens become superscript runs.
    """
    p = cell.paragraphs[0]
    # Remove all existing runs
    for run in list(p.runs):
        run._element.getparent().remove(run._element)

    parts = _SUPER_RE.split(text)
    # parts alternates: plain, number, plain, number, ...
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:  # odd indices are the captured group (the number)
            run.font.superscript = True


def set_row_text(row, values):
    for cell, val in zip(row.cells, values):
        set_cell_text(cell, val)


def delete_table(doc, table):
    """Remove a table element from the document body."""
    tbl = table._element
    tbl.getparent().remove(tbl)


def main():
    doc = Document(str(SRC))
    table = doc.tables[0]  # Table 0 = the empty template

    # Walk rows; fill subheadings with our entries
    rows_snapshot = list(table.rows)

    section_map = {}
    heading_rows = {}
    current_heading = None
    for row in rows_snapshot:
        txts = [c.text.strip() for c in row.cells]
        if txts[0] == "REAGENT or RESOURCE":
            continue
        if is_subheading_row(row):
            current_heading = txts[0]
            heading_rows[current_heading] = row
            section_map[current_heading] = []
        else:
            if current_heading is not None:
                section_map[current_heading].append(row)

    for heading in SUBHEADINGS_IN_TEMPLATE:
        empty_rows = section_map.get(heading, [])
        entries = KRT.get(heading, [])

        if not entries:
            for er in empty_rows:
                remove_row(table, er)
            if heading in heading_rows:
                remove_row(table, heading_rows[heading])
            continue

        for i, entry in enumerate(entries):
            if i < len(empty_rows):
                set_row_text(empty_rows[i], entry)
            else:
                ref = empty_rows[-1] if empty_rows else None
                if ref is None:
                    print(f"WARN: no template rows under {heading}")
                    continue
                new_row = insert_row_after(table, empty_rows[-1] if i > 0 else empty_rows[0])
                set_row_text(new_row, entry)
                empty_rows.append(new_row)

        for j in range(len(entries), len(empty_rows)):
            remove_row(table, empty_rows[j])

    # Clean up any stray rows
    for row in list(table.rows):
        txts = [c.text.strip() for c in row.cells]
        if txts[0] == "REAGENT or RESOURCE":
            continue
        if is_subheading_row(row) and txts[0] not in KRT:
            remove_row(table, row)
            continue
        if all(t == "" for t in txts):
            remove_row(table, row)

    # Delete Appendix A example tables (Tables 1 and 2 in the template).
    # Iterate in reverse so indices stay valid after each deletion.
    for t in reversed(doc.tables[1:]):
        delete_table(doc, t)

    doc.save(str(DST))
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
